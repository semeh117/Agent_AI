"""
send_results_email.py
--------------
Creates a Gmail DRAFT (never sends automatically) addressed to the
candidate, containing their ranked job matches and the generated cover
letter for the top match.

Why a draft, not a send: the Gmail scope used here (gmail.compose) only
supports creating drafts, not sending — and honestly, that's the right
default anyway. A job application email going out without the candidate
looking at it first is not something this project should automate fully.

Requires credentials.json (OAuth client, Desktop app type) in the
project root — see project setup notes for how to generate one via
Google Cloud Console. On first run, this opens a browser window for you
to authorize; after that, a token.json is saved so you won't need to
re-authorize every time (until the token expires or is revoked).

credentials.json and token.json are both secrets — already covered by
.gitignore, never commit them.
"""

import base64
import io
import os
import re
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

from docx import Document

from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build

# Deliberately narrow scope: lets this code create drafts only. It
# CANNOT read your inbox, send mail, or delete anything with this scope.
SCOPES = ["https://www.googleapis.com/auth/gmail.compose"]

CREDENTIALS_PATH = "credentials.json"
TOKEN_PATH = "token.json"


def _get_gmail_service():
    """
    Handles the OAuth flow: reuses a saved token.json if present and
    still valid, refreshes it if expired, or runs the browser-based
    authorization flow (using credentials.json) if this is the first
    run. Returns an authenticated Gmail API client.
    """
    creds = None
    if os.path.exists(TOKEN_PATH):
        creds = Credentials.from_authorized_user_file(TOKEN_PATH, SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists(CREDENTIALS_PATH):
                raise FileNotFoundError(
                    f"{CREDENTIALS_PATH} not found in project root. "
                    "Download it from Google Cloud Console > Credentials "
                    "(OAuth client, Desktop app type) and place it here."
                )
            flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_PATH, SCOPES)
            creds = flow.run_local_server(port=0)  # opens a browser window

        with open(TOKEN_PATH, "w", encoding="utf-8") as token_file:
            token_file.write(creds.to_json())

    return build("gmail", "v1", credentials=creds)


def _build_email_body(cv_info, ranked_jobs: list, cover_letter: str) -> str:
    """
    Formats the ranked jobs + cover letter into one plain-text email body.
    """
    lines = [
        f"Hi {cv_info.full_name or ''},".strip(),
        "",
        "Here are your job matches from this run of the AI Job Matching Assistant:",
        "",
    ]

    for i, job in enumerate(ranked_jobs, start=1):
        line = f"{i}. {job['job_title']} @ {job['company']} — {job['score_percent']}% match"
        if job.get("inconclusive"):
            line += " (inconclusive — no extractable requirements, score is not a real skill match)"
        lines.append(line)
        matching = [m["job_skill"] for m in job["skills_detail"]["matching"]]
        missing = job["skills_detail"]["missing"]
        if matching:
            lines.append(f"   Matching: {', '.join(matching)}")
        if missing:
            lines.append(f"   Missing:  {', '.join(missing)}")
        lines.append(f"   {job['url']}")
        lines.append("")

    if ranked_jobs:
        top = ranked_jobs[0]
        lines.append("=" * 60)
        lines.append(f"Cover letter for your top match — {top['job_title']} @ {top['company']}:")
        lines.append("=" * 60)
        lines.append("")
        lines.append(cover_letter)

    return "\n".join(lines)


def _build_cover_letter_docx(cv_info, cover_letter: str, job_title: str, company: str) -> bytes:
    """
    Renders the cover letter as a Word (.docx) document, returned as bytes
    in memory so it can be attached to the draft without touching disk.
    """
    doc = Document()

    heading = doc.add_heading("Cover Letter", level=0)
    heading.alignment = 1  # CENTER

    if cv_info.full_name:
        name_para = doc.add_paragraph(cv_info.full_name)
        name_para.alignment = 1  # CENTER

    doc.add_paragraph(f"Position: {job_title}")
    doc.add_paragraph(f"Company: {company}")
    doc.add_paragraph()

    for para in cover_letter.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _safe_attachment_name(job_title: str, company: str) -> str:
    """Windows/email-safe filename derived from the job title + company."""
    base = re.sub(r'[\\/:*?"<>|]+', "_", f"Cover Letter - {job_title} @ {company}")
    return base.strip()[:80] + ".docx"


def create_results_draft(cv_info, ranked_jobs: list, cover_letter: str, to_email: str) -> dict:
    """
    Creates a Gmail draft addressed to to_email, containing the ranked
    job matches and the cover letter for the top match. Does NOT send —
    the candidate opens Gmail and reviews/sends it themselves.

    Args:
        cv_info: the candidate's CVInfo object.
        ranked_jobs: results from run_job_matching() — already sorted.
        cover_letter: output of generate_cover_letter() for the top job.
        to_email: the candidate's own email address (self-addressed).

    Returns:
        The Gmail API's draft resource dict (includes the draft "id"),
        so the caller can confirm it was created and, if wanted, print
        a direct link to it.
    """
    service = _get_gmail_service()

    body_text = _build_email_body(cv_info, ranked_jobs, cover_letter)
    subject = f"Your job matches — {ranked_jobs[0]['job_title']} and {len(ranked_jobs) - 1} more" \
        if ranked_jobs else "Your job matching results"

    message = MIMEMultipart("mixed")
    message["to"] = to_email
    message["subject"] = subject

    # Plain-text body: ranked jobs + the cover letter text.
    message.attach(MIMEText(body_text))

    # Attach the cover letter as a Word document so the candidate can open,
    # edit, and submit it directly.
    if cover_letter:
        top = ranked_jobs[0] if ranked_jobs else {"job_title": "Job", "company": ""}
        docx_bytes = _build_cover_letter_docx(cv_info, cover_letter, top["job_title"], top["company"])
        attachment = MIMEApplication(
            docx_bytes,
            _subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        attachment.add_header(
            "Content-Disposition",
            "attachment",
            filename=_safe_attachment_name(top["job_title"], top["company"]),
        )
        message.attach(attachment)

    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode("utf-8")

    draft = service.users().drafts().create(
        userId="me",
        body={"message": {"raw": raw_message}},
    ).execute()

    print(f"Draft created (id: {draft['id']}). Open Gmail > Drafts to review and send.")
    return draft